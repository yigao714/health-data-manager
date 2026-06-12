"""
健康数据分析报告生成器 — Word (.docx) 格式
==========================================

从 HealthDataStore 生成格式化的 Word 文档报告。
包含个人档案、健康评分、趋势图表、风险评估和完整数据表。
"""

import io
import math
from pathlib import Path
from datetime import datetime
from typing import Optional

import matplotlib
matplotlib.use('Agg')  # 无头模式
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from health_data_schema import HealthDataStore


# ── 配色 ──
C_BG = '#111827'
C_BLUE = '#6366f1'
C_CYAN = '#22d3ee'
C_GREEN = '#34d399'
C_ORANGE = '#fb923c'
C_RED = '#f87171'
C_TEXT = '#e2e8f0'
C_GRID = '#1e293b'


def _avg(values):
    valid = [v for v in values if v is not None]
    return sum(valid) / len(valid) if valid else 0


def _calc_scores(records, profile):
    """计算健康评分（与 dashboard.html calcScores 逻辑一致）"""
    if not records:
        return {}

    avg_steps = _avg([r.steps for r in records])
    avg_exercise = _avg([r.exercise_minutes for r in records])
    avg_rest_hr = _avg([r.resting_heart_rate for r in records])
    avg_sleep = _avg([r.sleep_hours for r in records])
    avg_sleep_score = _avg([r.sleep_score for r in records])

    # 运动活力
    step_score = min(100, round(avg_steps / 7000 * 100))
    exercise_score = min(100, round(avg_exercise / (150/7) * 100))
    act_score = round(step_score * 0.7 + exercise_score * 0.3)

    # 心脏健康
    optimal_hr = 65
    hr_tolerance = 3
    if profile.age:
        if profile.age >= 65:
            optimal_hr = 70; hr_tolerance = 2
        elif profile.age >= 50:
            optimal_hr = 68; hr_tolerance = 2.5

    is_beta_blocker = profile.has_beta_blocker
    if is_beta_blocker and avg_rest_hr < 60:
        heart_score = 80
    else:
        hr_dev = abs(avg_rest_hr - optimal_hr)
        heart_score = max(0, min(100, round(100 - hr_dev * hr_tolerance)))

    # 睡眠质量
    rec_min = 6 if (profile.age and profile.age >= 65) else 7
    rec_max = 8 if (profile.age and profile.age >= 65) else 9

    if rec_min <= avg_sleep <= rec_max:
        dur_score = 100
    elif avg_sleep >= rec_min - 1:
        dur_score = 70
    elif avg_sleep >= rec_min - 2:
        dur_score = 40
    else:
        dur_score = 20

    # SRI
    sleep_starts = []
    for r in records:
        if r.sleep_start:
            parts = r.sleep_start.split(':')
            h, m = int(parts[0]), int(parts[1])
            hrs = h + m/60
            if hrs < 12:
                hrs += 24
            sleep_starts.append(hrs)

    if len(sleep_starts) >= 2:
        sri_std = float(np.std(sleep_starts)) * 60
        sri_score = 100 if sri_std < 30 else 75 if sri_std < 60 else 50 if sri_std < 90 else 25
    else:
        sri_score = 50

    sleep_score = round(dur_score * 0.3 + avg_sleep_score * 0.5 + sri_score * 0.2)

    # 综合
    overall = round(act_score * 0.3 + heart_score * 0.3 + sleep_score * 0.4)

    # 血氧（使用集中阈值 + 精度容差）
    thr = profile.get_thresholds()
    eff_danger = thr['spo2_danger'] - thr['spo2_precision']
    eff_warning = thr['spo2_warning'] - thr['spo2_precision']
    
    has_spo2 = any(r.spo2_min is not None for r in records)
    if not has_spo2:
        spo2_text = '无血氧监测数据'
        spo2_level = '未监测'
    else:
        spo2_danger = sum(1 for r in records if r.spo2_min is not None and r.spo2_min < eff_danger)
        spo2_warn = sum(1 for r in records if r.spo2_min is not None and r.spo2_min < eff_warning)
        if spo2_danger > 0:
            spo2_text = f'{spo2_danger}天SpO2<{eff_danger}%'
            spo2_level = '需关注'
        elif spo2_warn > len(records) * 0.3:
            spo2_text = f'{spo2_warn}天SpO2<{eff_warning}%'
            spo2_level = '留意'
        else:
            spo2_text = '未见低氧事件'
            spo2_level = '正常'

    return {
        'act_score': act_score, 'heart_score': heart_score,
        'sleep_score': sleep_score, 'overall': overall,
        'spo2_level': spo2_level, 'spo2_text': spo2_text,
        'avg_steps': round(avg_steps), 'avg_exercise': round(avg_exercise, 1),
        'avg_rest_hr': round(avg_rest_hr, 1), 'avg_sleep': round(avg_sleep, 2),
        'avg_sleep_score': round(avg_sleep_score, 1),
    }


def _generate_chart(records, field, ylabel, color, title, goal_line=None):
    """生成单指标趋势图，返回 PNG bytes"""
    dates = [r.date for r in records]
    values = [getattr(r, field) for r in records]

    fig, ax = plt.subplots(figsize=(8, 3), facecolor=C_BG)
    ax.set_facecolor(C_BG)

    # 过滤 None
    valid = [(d, v) for d, v in zip(dates, values) if v is not None]
    if not valid:
        plt.close(fig)
        return None

    ds, vs = zip(*valid)
    ax.plot(ds, vs, color=color, linewidth=2, marker='o', markersize=4)
    ax.fill_between(ds, vs, alpha=0.15, color=color)

    if goal_line is not None:
        ax.axhline(y=goal_line, color=C_GREEN, linestyle='--', alpha=0.5, linewidth=1)

    ax.set_title(title, color=C_TEXT, fontsize=12, pad=10)
    ax.set_ylabel(ylabel, color=C_TEXT, fontsize=9)
    ax.tick_params(colors=C_TEXT, labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color(C_GRID)
    ax.spines['left'].set_color(C_GRID)
    ax.grid(axis='y', color=C_GRID, alpha=0.3)

    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, facecolor=C_BG)
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)


def _add_paragraph(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(11)
    return p


def _score_color(score):
    if score >= 80:
        return RGBColor(0x34, 0xd3, 0x99)
    elif score >= 60:
        return RGBColor(0xfb, 0x92, 0x3c)
    else:
        return RGBColor(0xf8, 0x71, 0x71)


def generate_docx_report(store: HealthDataStore, output_path: str):
    """生成 Word 健康报告"""
    # 设置中文字体
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

    doc = Document()

    records = store.records
    profile = store.profile
    scores = _calc_scores(records, profile)

    if not records:
        doc.add_paragraph("暂无数据")
        doc.save(output_path)
        return

    # ═══ 封面 ═══
    doc.add_paragraph()
    title = doc.add_heading('健康数据分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n{store.person_name or "未命名"}')
    run.font.size = Pt(24)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n数据范围: {records[0].date} ~ {records[-1].date}')
    run.font.size = Pt(12)
    run = p.add_run(f'\n共 {len(records)} 天数据')
    run.font.size = Pt(12)
    run = p.add_run(f'\n生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_page_break()

    # ═══ 个人档案 + 综合评分 ═══
    _add_heading(doc, '个人健康档案', level=1)

    if profile.age or profile.sex or profile.height_cm:
        items = []
        if profile.age:
            items.append(f'年龄: {profile.age}岁')
        if profile.sex:
            items.append(f'性别: {"男" if profile.sex == "male" else "女"}')
        if profile.height_cm and profile.weight_kg:
            bmi = round(profile.weight_kg / (profile.height_cm/100)**2, 1)
            cat = '偏瘦' if bmi < 18.5 else '正常' if bmi < 24 else '超重' if bmi < 28 else '肥胖'
            items.append(f'身高: {profile.height_cm}cm  体重: {profile.weight_kg}kg  BMI: {bmi} ({cat})')
        if profile.conditions:
            items.append(f'基础疾病: {", ".join(profile.conditions)}')
        if profile.medications:
            items.append(f'用药: {", ".join(profile.medications)}')

        for item in items:
            _add_paragraph(doc, f'  • {item}')
    else:
        _add_paragraph(doc, '（未填写个人档案）')

    doc.add_paragraph()
    _add_heading(doc, '综合健康评分', level=1)

    # 评分表格
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    headers = ['维度', '评分', '参考标准']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data_rows = [
        ('🏃 运动活力', scores['act_score'], 'Lancet 2022 + WHO 2020'),
        ('❤️ 心脏健康', scores['heart_score'], 'AHA + Copenhagen Study'),
        ('😴 睡眠质量', scores['sleep_score'], 'NSF + SRI'),
        ('⭐ 综合评分', scores['overall'], '三维加权'),
    ]
    for i, (dim, score, ref) in enumerate(data_rows):
        table.rows[i+1].cells[0].text = dim
        table.rows[i+1].cells[1].text = f'{score}分'
        table.rows[i+1].cells[2].text = ref

    p = doc.add_paragraph()
    run = p.add_run(f'\n🫁 血氧监测: {scores["spo2_level"]} — {scores["spo2_text"]}')
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('（血氧不参与评分，腕式传感器精度±2-3%，仅作参考标记）')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_page_break()

    # ═══ 趋势图表 ═══
    _add_heading(doc, '数据趋势', level=1)

    charts = [
        ('steps', '步数', C_BLUE, '每日步数', 7000),
        ('resting_heart_rate', 'bpm', C_RED, '静息心率', None),
        ('sleep_hours', '小时', C_CYAN, '睡眠时长', 7),
        ('sleep_score', '分', '#a78bfa', '睡眠分数', None),
    ]

    for field, ylabel, color, title, goal in charts:
        buf = _generate_chart(records, field, ylabel, color, title, goal)
        if buf:
            doc.add_picture(buf, width=Inches(6))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ═══ 统计摘要 ═══
    _add_heading(doc, '统计摘要', level=1)

    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Light Grid Accent 1'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_data = [
        ('指标', '均值'),
        ('日均步数', f'{scores["avg_steps"]} 步'),
        ('日均锻炼', f'{scores["avg_exercise"]} 分钟'),
        ('平均静息心率', f'{scores["avg_rest_hr"]} bpm'),
        ('平均睡眠时长', f'{scores["avg_sleep"]} 小时'),
        ('平均睡眠分数', f'{scores["avg_sleep_score"]} 分'),
    ]
    for i, (k, v) in enumerate(summary_data):
        table2.rows[i].cells[0].text = k
        table2.rows[i].cells[1].text = v

    doc.add_page_break()

    # ═══ 完整数据表 ═══
    _add_heading(doc, '完整数据', level=1)

    cols = ['日期', '步数', '距离km', '锻炼min', '静息心率', '血氧最低', '睡眠h', '睡眠分数']
    table3 = doc.add_table(rows=len(records)+1, cols=len(cols))
    table3.style = 'Light Grid Accent 1'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, c in enumerate(cols):
        table3.rows[0].cells[i].text = c

    for i, r in enumerate(records):
        row = table3.rows[i+1]
        row.cells[0].text = r.date
        row.cells[1].text = str(r.steps or '-')
        row.cells[2].text = str(r.distance_km or '-')
        row.cells[3].text = str(r.exercise_minutes or '-')
        row.cells[4].text = str(r.resting_heart_rate or '-')
        row.cells[5].text = f'{r.spo2_min}%' if r.spo2_min else '-'
        row.cells[6].text = str(r.sleep_hours or '-')
        row.cells[7].text = str(r.sleep_score or '-')

    # 缩小表格字体
    for row in table3.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()

    # ═══ 附录 ═══
    _add_heading(doc, '附录: 评估标准来源', level=1)

    refs = [
        'Lancet 2022: Paluch AE, et al. "Daily steps and all-cause mortality." Lancet Public Health 2022;7(3):e219-e228.',
        'WHO 2020: WHO guidelines on physical activity and sedentary behaviour. Geneva: WHO, 2020.',
        'AHA: American Heart Association — 成人正常静息心率 60-100 bpm.',
        'Copenhagen City Heart Study: 静息心率与心血管死亡风险独立相关.',
        '血氧: 腕式血氧参考(精度±2-3%)，非诊断。阈值因COPD等基础疾病自适应调整(参考BTS/GOLD)。',
        'NSF: Hirshkowitz M, et al. "Sleep time duration recommendations." Sleep Health 2015;1(1):40-43.',
        'SRI: Lunsford-Avery JR, et al. "Sleep/Wake Regularity and Cardiometabolic Risk." Sleep 2018;41(6).',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(f'• {ref}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    p = doc.add_paragraph()
    run = p.add_run('\n⚠️ 本报告由 AI 辅助生成，评分仅供参考，不构成医疗建议。')
    run.font.size = Pt(10)
    run.bold = True

    doc.save(output_path)
    print(f"  ✅ Word 报告已生成: {output_path}")
